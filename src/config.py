"""JobHunter 配置管理模块。

负责 config.yaml 的读取、写入与默认值管理。
"""
from __future__ import annotations

import os
import shutil
from pathlib import Path
from typing import Any, Dict, Optional

import yaml

# 技能包根目录（src/ 的上一级）
SKILL_ROOT = Path(__file__).resolve().parent.parent
CONFIG_PATH = SKILL_ROOT / "config.yaml"
DATA_DIR = SKILL_ROOT / "data"
EXPORT_DIR = SKILL_ROOT / "export"
LOGS_DIR = SKILL_ROOT / "logs"

DEFAULT_CONFIG: Dict[str, Any] = {
    "search": {
        "keywords": [],
        "city": "",
        "salary_min": 0,
        "salary_max": 0,
        "education": [],       # 学历过滤：如 ["本科", "大专"]
        "experience": [],      # 经验过滤：如 ["应届毕业生", "在校生", "无需经验"]
        "veto_words": [],
        "city_filter_enabled": True,   # 按目标城市过滤开关：采集后仅保留目标城市(region)命中岗位
        "target_cities": ["厦门", "福州", "泉州"],  # 目标城市名（子串匹配岗位真实所在地）
    },
    "resume": {
        "files": [],         # 简历文件路径列表（.docx / .pdf）
        "file": "",          # 兼容旧版：单份简历文件路径
        "info": "",          # 简历信息库（结构化文本）
    },
    "llm": {
        "api_key": "ollama",
        "base_url": "http://localhost:11434/v1",
        "model": "qwen2.5:7b",
        "temperature": 0.3,
    },
    "monitor": {
        "enabled": False,
        "interval_minutes": 30,
    },
    "browser": {
        "headless": False,
        "send_interval_min": 30,
        "send_interval_max": 60,
    },
    "platforms": {
        # 多平台采集配置：每个平台的独立浏览器 profile 目录与是否启用。
        # 来源标签由 platforms.py 统一管理；这里仅为每个平台补充 profile/开关。
        "boss": {"enabled": True},
        "yingsheng": {"enabled": True, "profile": ""},
        "zhaopin": {"enabled": True, "profile": ""},
        "51job": {"enabled": True, "profile": ""},
        "qiuzhifangzhou": {"enabled": True, "profile": ""},
    },
}


def ensure_dirs() -> None:
    """确保数据/导出/日志目录存在。"""
    for d in (DATA_DIR, EXPORT_DIR, LOGS_DIR):
        d.mkdir(parents=True, exist_ok=True)


def load_config() -> Dict[str, Any]:
    """读取配置，若不存在则返回默认配置。"""
    ensure_dirs()
    if not CONFIG_PATH.exists():
        return _deep_copy(DEFAULT_CONFIG)
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f) or {}
    # 合并默认值，保证缺失字段有兜底
    merged = _deep_copy(DEFAULT_CONFIG)
    _deep_merge(merged, data)
    return merged


def save_config(config: Dict[str, Any]) -> None:
    """写入配置到 config.yaml。"""
    ensure_dirs()
    with open(CONFIG_PATH, "w", encoding="utf-8") as f:
        yaml.safe_dump(config, f, allow_unicode=True, sort_keys=False)


def is_configured() -> bool:
    """判断是否已完成首次配置（存在 config.yaml 且有关键词）。"""
    if not CONFIG_PATH.exists():
        return False
    cfg = load_config()
    return bool(cfg.get("search", {}).get("keywords"))


def get_platform_cfg(key: str) -> Dict[str, Any]:
    """取某平台的采集配置（config.yaml platforms.<key>），不存在则返回空 dict。"""
    cfg = load_config()
    return (cfg.get("platforms") or {}).get(key, {}) or {}


def get_llm_client():
    """创建 OpenAI 兼容客户端（支持 Ollama / 硅基流动等）。"""
    from openai import OpenAI

    cfg = load_config()
    llm = cfg.get("llm", {})
    api_key = llm.get("api_key", "ollama")
    return OpenAI(
        api_key=api_key,
        base_url=llm.get("base_url", "http://localhost:11434/v1"),
    )


def _deep_copy(obj: Any) -> Any:
    import copy

    return copy.deepcopy(obj)


def _deep_merge(base: Dict[str, Any], override: Dict[str, Any]) -> None:
    """递归合并 override 到 base。"""
    for k, v in override.items():
        if k in base and isinstance(base[k], dict) and isinstance(v, dict):
            _deep_merge(base[k], v)
        else:
            base[k] = v


def copy_resume_to_skill(src_path: str) -> Optional[str]:
    """将用户上传的简历复制到技能包 data/ 目录，返回新路径。"""
    src = Path(src_path)
    if not src.exists():
        return None
    ensure_dirs()
    dest = DATA_DIR / f"resume{src.suffix}"
    shutil.copy2(src, dest)
    return str(dest)


def extract_resume_text(resume_path: str) -> str:
    """从简历文件提取纯文本，支持 .md / .txt / .docx / .pdf。"""
    path = Path(resume_path)
    if not path.exists():
        return ""
    suffix = path.suffix.lower()
    try:
        if suffix in (".md", ".txt"):
            return path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".docx":
            import docx

            doc = docx.Document(str(path))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        if suffix == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            parts = []
            for page in reader.pages:
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(text)
            return "\n".join(parts)
    except Exception as e:
        return f"（简历解析失败：{e}）"
    return ""


def load_resume_text() -> str:
    """读取简历素材（多份简历文件 + 信息库），返回合并后的纯文本。"""
    cfg = load_config()
    resume_cfg = cfg.get("resume", {})
    parts = []

    # 多份简历文件（优先）
    files = list(resume_cfg.get("files") or [])
    # 兼容旧版单文件字段
    legacy = resume_cfg.get("file", "")
    if legacy and Path(legacy).exists() and legacy not in files:
        files.append(legacy)

    for fp in files:
        if not Path(fp).exists():
            continue
        text = extract_resume_text(fp)
        if text and not text.startswith("（简历解析失败"):
            parts.append(f"===== 简历：{Path(fp).name} =====\n{text}")

    # 信息库
    info = resume_cfg.get("info", "")
    if info:
        parts.append(f"===== 简历信息库 =====\n{info}")

    if parts:
        return "\n\n".join(parts)
    return "（未提供简历素材，请先在配置面板上传简历或填写信息库）"
